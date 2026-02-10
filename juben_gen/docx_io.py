from __future__ import annotations

from dataclasses import dataclass
from pathlib import Path
from typing import Iterable, List, Optional

from docx import Document


def read_docx_lines(path: str | Path) -> List[str]:
    """
    读取 docx 的段落与表格文本，按“行”返回（空行剔除）。
    """
    p = Path(path)
    doc = Document(str(p))
    lines: List[str] = []

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        if text:
            lines.append(text)

    for table in doc.tables:
        for row in table.rows:
            cells = [(c.text or "").strip() for c in row.cells]
            if any(cells):
                lines.append(" | ".join(cells))

    return lines


def write_docx_lines(
    path: str | Path,
    lines: Iterable[str],
    *,
    title: Optional[str] = None,
) -> None:
    """
    将“行”写入 docx：每行一个段落。
    """
    p = Path(path)
    p.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    if title:
        doc.add_heading(title, level=1)

    for line in lines:
        t = (line or "").rstrip()
        if t:
            doc.add_paragraph(t)
        else:
            doc.add_paragraph("")

    doc.save(str(p))


@dataclass(frozen=True)
class DocxSnippet:
    """
    用于 prompt few-shot：只放少量示例行，避免超长。
    """

    title: str
    lines: List[str]


def load_docx_snippet(path: str | Path, *, title: str, max_lines: int = 60) -> DocxSnippet:
    lines = read_docx_lines(path)
    return DocxSnippet(title=title, lines=lines[:max_lines])

